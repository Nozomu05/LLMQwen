import os
import zipfile
import shutil
import tempfile
from pathlib import Path
from dotenv import load_dotenv
from typing import List
import re
import csv
import xml.etree.ElementTree as ET

from langchain_community.document_loaders import (
    DirectoryLoader, 
    TextLoader, 
    UnstructuredPowerPointLoader, 
    Docx2txtLoader, 
    UnstructuredODTLoader,
    PyMuPDFLoader,
    UnstructuredPDFLoader,
    UnstructuredMarkdownLoader,
    UnstructuredWordDocumentLoader
)
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.embeddings import FastEmbedEmbeddings
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_chroma import Chroma
from langchain_core.documents import Document
from concurrent.futures import ThreadPoolExecutor, as_completed
import logging
from datetime import datetime

try:
    from docx import Document as DocxDocument
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False

try:
    from docx import Document as DocxDocument
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False

def load_docx_with_python_docx(file_path: str) -> List[Document]:
    if not PYTHON_DOCX_AVAILABLE:
        raise ImportError("python-docx not available")
    doc = DocxDocument(file_path)
    text = "\n\n".join([para.text for para in doc.paragraphs if para.text.strip()])
    if not text.strip():
        raise ValueError("No text extracted")
    return [Document(page_content=text, metadata={"source": Path(file_path).name})]

def load_docx_raw_xml(file_path: str) -> List[Document]:
    with zipfile.ZipFile(file_path, 'r') as docx_zip:
        try:
            xml_content = docx_zip.read('word/document.xml')
            root = ET.fromstring(xml_content)
            
            namespaces = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
            paragraphs = root.findall('.//w:t', namespaces)
            text = "\n".join([p.text for p in paragraphs if p.text])
            
            if not text.strip():
                raise ValueError("No text extracted from XML")
            
            return [Document(page_content=text, metadata={"source": Path(file_path).name})]
        except Exception as e:
            raise ValueError(f"Failed to extract from XML: {e}")


def _format_docx_table(table) -> str:
    """Format a python-docx Table object as a markdown table string."""
    if not table.rows:
        return ""
    rows_text = []
    for row in table.rows:
        cells = [cell.text.replace('\n', ' ').strip() for cell in row.cells]
        rows_text.append("| " + " | ".join(cells) + " |")
    if not rows_text:
        return ""
    n_cols = len(table.rows[0].cells)
    separator = "| " + " | ".join(["---"] * n_cols) + " |"
    return "\n".join([rows_text[0], separator] + rows_text[1:])


def load_docx_table_aware(file_path: str) -> List[Document]:
    """Load DOCX preserving table structure as separate markdown Documents.

    Iterates body elements in document order:
    - Consecutive paragraphs → one text Document (type='text')
    - Each table → one markdown Document (type='table')
    This prevents BD-rate tables from being merged with surrounding text
    and then split at arbitrary character boundaries.
    """
    if not PYTHON_DOCX_AVAILABLE:
        raise ImportError("python-docx not available")

    from docx.table import Table as DocxTable
    from docx.text.paragraph import Paragraph as DocxParagraph

    doc = DocxDocument(file_path)
    source = Path(file_path).name
    result_docs: List[Document] = []
    current_paragraphs: List[str] = []

    def _flush_paragraphs():
        nonlocal current_paragraphs
        text = "\n\n".join(p for p in current_paragraphs if p)
        if text.strip():
            result_docs.append(Document(
                page_content=text,
                metadata={"source": source, "type": "text"},
            ))
        current_paragraphs = []

    for element in doc.element.body:
        tag = element.tag.split('}')[-1] if '}' in element.tag else element.tag
        if tag == 'tbl':
            _flush_paragraphs()
            table = DocxTable(element, doc)
            md = _format_docx_table(table)
            if md.strip():
                result_docs.append(Document(
                    page_content=md,
                    metadata={"source": source, "type": "table"},
                ))
        elif tag == 'p':
            para = DocxParagraph(element, doc)
            text = para.text.strip()
            if text:
                current_paragraphs.append(text)

    _flush_paragraphs()

    if not result_docs:
        raise ValueError("No content extracted by table-aware loader")
    return result_docs


class TableAwareSplitter:
    """Splits documents while preserving markdown table structure.

    - type='table' docs: split at row boundaries, prepending header+separator
      to every chunk so each chunk is a self-contained table.
    - all other docs: delegate to RecursiveCharacterTextSplitter.
    """

    def __init__(self, chunk_size: int, chunk_overlap: int):
        self._chunk_size = chunk_size
        self._text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            separators=["\n\n", "\n", ". ", " ", ""],
        )

    def split_documents(self, docs: List[Document]) -> List[Document]:
        result: List[Document] = []
        for doc in docs:
            if doc.metadata.get("type") == "table":
                result.extend(self._split_table(doc))
            else:
                result.extend(self._text_splitter.split_documents([doc]))
        return result

    def _split_table(self, doc: Document) -> List[Document]:
        lines = doc.page_content.split('\n')
        if len(lines) < 3:
            return [doc]
        header = lines[0]
        sep_line = lines[1]
        data_rows = lines[2:]
        header_prefix = header + '\n' + sep_line + '\n'

        chunks: List[Document] = []
        current_rows: List[str] = []

        for row in data_rows:
            candidate = header_prefix + '\n'.join(current_rows + [row])
            if len(candidate) > self._chunk_size and current_rows:
                chunk_text = header_prefix + '\n'.join(current_rows)
                chunks.append(Document(page_content=chunk_text, metadata=doc.metadata))
                current_rows = [row]
            else:
                current_rows.append(row)

        if current_rows:
            chunk_text = header_prefix + '\n'.join(current_rows)
            chunks.append(Document(page_content=chunk_text, metadata=doc.metadata))

        return chunks if chunks else [doc]


class SemanticTableAwareSplitter:
    """Semantic chunking for text, table-aware splitting for tables.

    Uses SemanticChunker (embedding-based boundary detection) for prose/text
    sections, and row-boundary splitting for markdown tables so BD-rate data
    is never cut mid-row.

    breakpoint_threshold_type='gradient' finds where semantic similarity drops
    most sharply — best for technical documents with varied section lengths.
    """

    def __init__(self, embeddings, chunk_size: int, chunk_overlap: int):
        from langchain_experimental.text_splitter import SemanticChunker
        self._semantic = SemanticChunker(
            embeddings,
            breakpoint_threshold_type="gradient",
        )
        self._table_splitter = TableAwareSplitter(chunk_size, chunk_overlap)

    def split_documents(self, docs: List[Document]) -> List[Document]:
        text_docs = [d for d in docs if d.metadata.get("type") != "table"]
        table_docs = [d for d in docs if d.metadata.get("type") == "table"]
        result: List[Document] = []
        if text_docs:
            result.extend(self._semantic.split_documents(text_docs))
        if table_docs:
            result.extend(self._table_splitter.split_documents(table_docs))
        return result


# Regex matching MPEG/ISO numbered section headings and markdown headings:
#   "1  Introduction", "1.2.3 Method", "## Heading", "### Sub"
_SECTION_HEADING_RE = re.compile(
    r'(?m)^(?=(?:\d+\.)*\d+\s+\S|#{1,6}\s)',
)


class AdaptiveChunker:
    """Structure-first adaptive chunking — the most complete strategy.

    Pipeline per text document:
      1. Split on section headings (numbered MPEG/ISO sections + markdown headings)
         so each chunk respects the document's own organisation.
      2. Sections that are still too long → apply SemanticChunker (gradient)
         to find the best internal split point by meaning.
      3. Sections that are very short → merge with adjacent sections until
         they reach a minimum viable size (chunk_size // 4).
      4. Tables → row-boundary splitting (unchanged).

    This is the right choice for MPEG technical documents: they have clear
    numbered sections, varying lengths, BD-rate tables, and dense technical prose.
    """

    _MIN_SECTION = None  # set in __init__

    def __init__(self, embeddings, chunk_size: int, chunk_overlap: int):
        from langchain_experimental.text_splitter import SemanticChunker
        self._chunk_size = chunk_size
        self._min_section = chunk_size // 4
        self._semantic = SemanticChunker(
            embeddings,
            breakpoint_threshold_type="gradient",
        )
        self._table_splitter = TableAwareSplitter(chunk_size, chunk_overlap)

    # ------------------------------------------------------------------
    def split_documents(self, docs: List[Document]) -> List[Document]:
        text_docs = [d for d in docs if d.metadata.get("type") != "table"]
        table_docs = [d for d in docs if d.metadata.get("type") == "table"]
        result: List[Document] = []
        for doc in text_docs:
            result.extend(self._split_text(doc))
        result.extend(self._table_splitter.split_documents(table_docs))
        return result

    def _split_text(self, doc: Document) -> List[Document]:
        # Step 1 — split on section headings
        sections = self._split_on_headings(doc)

        # Step 2 — merge sections that are too short into their neighbour
        sections = self._merge_short_sections(sections, doc.metadata)

        # Step 3 — semantically split sections that are still too long
        result: List[Document] = []
        for sec in sections:
            if len(sec) <= self._chunk_size:
                result.append(Document(page_content=sec, metadata=doc.metadata))
            else:
                sub = self._semantic.split_documents(
                    [Document(page_content=sec, metadata=doc.metadata)]
                )
                result.extend(sub)
        return result

    def _split_on_headings(self, doc: Document) -> List[str]:
        text = doc.page_content
        boundaries = [m.start() for m in _SECTION_HEADING_RE.finditer(text)]
        if not boundaries:
            return [text]
        # include text before the first heading as its own section
        sections: List[str] = []
        starts = ([0] if boundaries[0] > 0 else []) + boundaries
        for i, start in enumerate(starts):
            end = starts[i + 1] if i + 1 < len(starts) else len(text)
            section = text[start:end].strip()
            if section:
                sections.append(section)
        return sections if sections else [text]

    def _merge_short_sections(self, sections: List[str], metadata: dict) -> List[str]:
        if not sections:
            return sections
        merged: List[str] = [sections[0]]
        for sec in sections[1:]:
            if len(merged[-1]) < self._min_section:
                merged[-1] = merged[-1] + '\n\n' + sec
            else:
                merged.append(sec)
        return merged


def load_excel(file_path: str) -> List[Document]:
    """Load .xlsx/.xlsm files using openpyxl, one Document per sheet."""
    import openpyxl
    wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
    docs = []
    source = Path(file_path).name
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        if not hasattr(ws, 'iter_rows'):  
            continue
        rows = []
        for row in ws.iter_rows(values_only=True):
            # Skip fully empty rows
            if all(cell is None for cell in row):
                continue
            rows.append(" | ".join("" if cell is None else str(cell) for cell in row))
        if not rows:
            continue
        text = f"[File: {source}] [Sheet: {sheet_name}]\n" + "\n".join(rows)
        docs.append(Document(page_content=text, metadata={"source": source}))
    wb.close()
    if not docs:
        raise ValueError("No data found in any sheet")
    return docs


def load_csv(file_path: str) -> List[Document]:
    """Load .csv files, converting tabular content to text."""
    source = Path(file_path).name
    for encoding in ("utf-8", "latin-1", "cp1252"):
        try:
            with open(file_path, newline="", encoding=encoding) as f:
                reader = csv.reader(f)
                rows = [" | ".join(cell.strip() for cell in row) for row in reader if any(c.strip() for c in row)]
            if not rows:
                raise ValueError("Empty CSV")
            text = f"[File: {source}]\n" + "\n".join(rows)
            return [Document(page_content=text, metadata={"source": source})]
        except (UnicodeDecodeError, csv.Error):
            continue
    raise ValueError(f"Could not decode CSV: {file_path}")


def extract_zip_files(docs_dir: Path) -> None:
    zip_files = list(docs_dir.glob("**/*.zip"))
    
    if not zip_files:
        return
    
    print(f"\nFound {len(zip_files)} ZIP file(s) to extract...\n")
    
    supported_extensions = {'.pptx', '.pdf', '.docx', '.md', '.odt', '.txt',
                            '.yaml', '.yml', '.cfg', '.h', '.hpp', '.c', '.cpp',
                            '.sh', '.pl', '.pm', '.tex', '.cmake', '.log'}
    processed_zips = set()
    
    def extract_recursive(zip_path: Path, level: int = 0):
        if str(zip_path) in processed_zips:
            return
        processed_zips.add(str(zip_path))
        
        indent = "  " * level
        try:
            extract_dir = zip_path.parent / zip_path.stem
            
            if extract_dir.exists():
                print(f"{indent}Skipping (already extracted): {zip_path.name}")
                return
            
            print(f"{indent}Extracting: {zip_path.name}")
            
            with zipfile.ZipFile(zip_path, 'r') as zip_ref:
                file_list = zip_ref.namelist()
                
                supported_files = [f for f in file_list 
                                 if any(f.lower().endswith(ext) for ext in supported_extensions)]
                nested_zips = [f for f in file_list if f.lower().endswith('.zip')]
                
                if supported_files or nested_zips:
                    zip_ref.extractall(extract_dir)
                    
                    if supported_files:
                        print(f"{indent}  → Found {len(supported_files)} supported document(s)")
                    
                    if nested_zips:
                        print(f"{indent}  → Found {len(nested_zips)} nested ZIP file(s)")
                        for nested_zip in nested_zips:
                            nested_zip_path = extract_dir / nested_zip
                            if nested_zip_path.exists():
                                extract_recursive(nested_zip_path, level + 1)
                    
                    print(f"{indent}  -> Extracted to: {extract_dir.name}/")
                else:
                    print(f"{indent}  ! No supported documents found (skipping)")
                    
        except zipfile.BadZipFile:
            print(f"{indent}  ✗ Error: {zip_path.name} is not a valid ZIP file")
        except Exception as e:
            print(f"{indent}  ✗ Error extracting {zip_path.name}: {e}")
    
    for zip_path in zip_files:
        extract_recursive(zip_path)
    
    print()

def load_documents_batch(docs_dir: Path, batch_size: int = 50) -> tuple[List[Document], dict]:
    all_docs = []
    stats = {
        'successful': [],
        'failed': [],
        'total_by_type': {}
    }

    log_file = Path("ingestion_errors.log")
    logging.basicConfig(
        level=logging.INFO,
        format='%(asctime)s - %(levelname)s - %(message)s',
        handlers=[
            logging.FileHandler(log_file, mode='w', encoding='utf-8'),
            logging.StreamHandler()
        ]
    )
    logger = logging.getLogger(__name__)
    
    logger.info(f"Starting document ingestion from: {docs_dir}")
    
    def load_single_file(file_path: Path, loader_configs: list) -> tuple[str, list, str]:
        for config_name, loader_func in loader_configs:
            try:
                docs = loader_func(str(file_path))
                if isinstance(docs, list):
                    return file_path.name, docs, config_name
                else:
                    loaded_docs = docs.load()
                    return file_path.name, loaded_docs, config_name
            except Exception as e:
                logger.debug(f"  {config_name} failed for {file_path.name}: {str(e)}")
                continue
        
        return file_path.name, [], "FAILED"
    
    loaders_config = [
        ("**/*.md", [
            ("UnstructuredMarkdown", lambda p: UnstructuredMarkdownLoader(p)),
            ("TextLoader", lambda p: TextLoader(p))
        ], "Markdown"),
        
        ("**/*.txt", [
            ("TextLoader-utf8", lambda p: TextLoader(p, encoding='utf-8')),
            ("TextLoader-latin1", lambda p: TextLoader(p, encoding='latin-1')),
            ("TextLoader-auto", lambda p: TextLoader(p, autodetect_encoding=True))
        ], "Text"),
        
        ("**/*.pdf", [
            ("PyMuPDF", lambda p: PyMuPDFLoader(p)),
            ("Unstructured-Fast", lambda p: UnstructuredPDFLoader(p, mode="elements")),
            ("Unstructured-HiRes", lambda p: UnstructuredPDFLoader(p, mode="elements", strategy="hi_res"))
        ], "PDF"),
        
        ("**/*.pptx", [
            ("UnstructuredPowerPoint", lambda p: UnstructuredPowerPointLoader(p))
        ], "PowerPoint"),
        
        ("**/*.ppt", [
            ("UnstructuredPowerPoint", lambda p: UnstructuredPowerPointLoader(p))
        ], "PowerPoint (legacy)"),
        
        ("**/*.docx", [
            ("TableAwareDocx", lambda p: load_docx_table_aware(p)),
            ("Docx2txt", lambda p: Docx2txtLoader(p)),
            ("PythonDocx", lambda p: load_docx_with_python_docx(p)),
            ("UnstructuredWord-Fast", lambda p: UnstructuredWordDocumentLoader(p, mode="single")),
            ("UnstructuredWord-Elements", lambda p: UnstructuredWordDocumentLoader(p, mode="elements")),
            ("RawXML", lambda p: load_docx_raw_xml(p))
        ], "Word"),
        
        ("**/*.doc", [
            ("UnstructuredWord", lambda p: UnstructuredWordDocumentLoader(p)),
            ("Docx2txt", lambda p: Docx2txtLoader(p))
        ], "Word (legacy)"),
        
        ("**/*.odt", [
            ("UnstructuredODT", lambda p: UnstructuredODTLoader(p))
        ], "ODT"),

        ("**/*.xlsx", [
            ("OpenPyXL", lambda p: load_excel(p))
        ], "Excel"),

        ("**/*.xlsm", [
            ("OpenPyXL", lambda p: load_excel(p))
        ], "Excel (macro)"),

        ("**/*.csv", [
            ("CSV", lambda p: load_csv(p))
        ], "CSV"),

        ("**/*.yaml", [
            ("TextLoader-utf8", lambda p: TextLoader(p, encoding='utf-8')),
            ("TextLoader-auto", lambda p: TextLoader(p, autodetect_encoding=True))
        ], "YAML"),

        ("**/*.yml", [
            ("TextLoader-utf8", lambda p: TextLoader(p, encoding='utf-8')),
            ("TextLoader-auto", lambda p: TextLoader(p, autodetect_encoding=True))
        ], "YAML (yml)"),

        ("**/*.cfg", [
            ("TextLoader-utf8", lambda p: TextLoader(p, encoding='utf-8')),
            ("TextLoader-auto", lambda p: TextLoader(p, autodetect_encoding=True))
        ], "Config"),

        ("**/*.h", [
            ("TextLoader-utf8", lambda p: TextLoader(p, encoding='utf-8')),
            ("TextLoader-latin1", lambda p: TextLoader(p, encoding='latin-1'))
        ], "C/C++ Header"),

        ("**/*.hpp", [
            ("TextLoader-utf8", lambda p: TextLoader(p, encoding='utf-8')),
            ("TextLoader-latin1", lambda p: TextLoader(p, encoding='latin-1'))
        ], "C++ Header"),

        ("**/*.c", [
            ("TextLoader-utf8", lambda p: TextLoader(p, encoding='utf-8')),
            ("TextLoader-latin1", lambda p: TextLoader(p, encoding='latin-1'))
        ], "C Source"),

        ("**/*.cpp", [
            ("TextLoader-utf8", lambda p: TextLoader(p, encoding='utf-8')),
            ("TextLoader-latin1", lambda p: TextLoader(p, encoding='latin-1'))
        ], "C++ Source"),

        ("**/*.sh", [
            ("TextLoader-utf8", lambda p: TextLoader(p, encoding='utf-8')),
            ("TextLoader-auto", lambda p: TextLoader(p, autodetect_encoding=True))
        ], "Shell Script"),

        ("**/*.pl", [
            ("TextLoader-utf8", lambda p: TextLoader(p, encoding='utf-8')),
            ("TextLoader-auto", lambda p: TextLoader(p, autodetect_encoding=True))
        ], "Perl"),

        ("**/*.pm", [
            ("TextLoader-utf8", lambda p: TextLoader(p, encoding='utf-8')),
            ("TextLoader-auto", lambda p: TextLoader(p, autodetect_encoding=True))
        ], "Perl Module"),

        ("**/*.tex", [
            ("TextLoader-utf8", lambda p: TextLoader(p, encoding='utf-8')),
            ("TextLoader-auto", lambda p: TextLoader(p, autodetect_encoding=True))
        ], "LaTeX"),

        ("**/*.cmake", [
            ("TextLoader-utf8", lambda p: TextLoader(p, encoding='utf-8')),
            ("TextLoader-auto", lambda p: TextLoader(p, autodetect_encoding=True))
        ], "CMake"),

        ("**/*.log", [
            ("TextLoader-utf8", lambda p: TextLoader(p, encoding='utf-8')),
            ("TextLoader-latin1", lambda p: TextLoader(p, encoding='latin-1'))
        ], "Log"),
    ]
    
    for glob_pattern, loader_strategies, doc_type in loaders_config:
        logger.info(f"\n{'='*60}")
        logger.info(f"Processing {doc_type} files...")
        logger.info(f"{'='*60}")
        
        all_files = list(docs_dir.glob(glob_pattern))
        valid_files = [f for f in all_files if f.is_file() and not f.name.startswith(('~$', '._'))]
        
        if not valid_files:
            logger.info(f"  No {doc_type} files found")
            continue
            
        if len(valid_files) < len(all_files):
            skipped = len(all_files) - len(valid_files)
            logger.info(f"  ! Skipping {skipped} temp file(s)")
        
        logger.info(f"  Found {len(valid_files)} {doc_type} file(s)")
        
        successful_count = 0
        failed_files = []
        
        for file_path in valid_files:
            filename, docs, method = load_single_file(file_path, loader_strategies)
            
            if method != "FAILED" and docs:
                all_docs.extend(docs)
                successful_count += 1
                stats['successful'].append({
                    'file': str(file_path),
                    'type': doc_type,
                    'chunks': len(docs),
                    'method': method
                })
                logger.info(f"  + {filename} ({len(docs)} chunks) [{method}]")
            else:
                failed_files.append(filename)
                stats['failed'].append({
                    'file': str(file_path),
                    'type': doc_type,
                    'error': 'All loading strategies failed'
                })
                logger.error(f"  ✗ {filename} - ALL STRATEGIES FAILED")
        
        stats['total_by_type'][doc_type] = {
            'total': len(valid_files),
            'successful': successful_count,
            'failed': len(failed_files)
        }
        
        logger.info(f"\n  Summary: {successful_count}/{len(valid_files)} {doc_type} files loaded successfully")
        if failed_files:
            logger.warning(f"  Failed files: {', '.join(failed_files)}")
    
    return all_docs, stats

def main() -> None:
    load_dotenv()

    docs_dir = Path(os.environ["DOCS_DIR"]).resolve()
    batch_size = int(os.environ["BATCH_SIZE"])
    chunk_size = int(os.environ["CHUNK_SIZE"])
    chunk_overlap = int(os.environ["CHUNK_OVERLAP"])

    if not docs_dir.exists():
        docs_dir.mkdir(parents=True, exist_ok=True)
        print(f"Created empty docs directory at: {docs_dir}")
        print("Add .md, .txt, .pdf, .pptx, .docx, .odt files or .zip archives and re-run this command.")
        return

    supported_extensions = {'.md', '.txt', '.pdf', '.pptx', '.ppt', '.docx', '.doc', '.odt',
                            '.yaml', '.yml', '.cfg',
                            '.h', '.hpp', '.c', '.cpp', '.sh', '.pl', '.pm', '.tex', '.cmake', '.log'}
    total_doc_count = sum(1 for f in docs_dir.rglob("*") if f.is_file() and f.suffix.lower() in supported_extensions)
    
    print(f"Loading documents from: {docs_dir}")
    print(f"Batch size: {batch_size} | Chunk size: {chunk_size} | Overlap: {chunk_overlap}")
    print(f"\n📁 Found {total_doc_count} total documents to process\n")

    extract_zip_files(docs_dir)

    all_docs, stats = load_documents_batch(docs_dir, batch_size)

    print("\n" + "="*70)
    print("INGESTION SUMMARY")
    print("="*70)
    
    total_files = len(stats['successful']) + len(stats['failed'])
    print(f"\nTotal files processed: {total_files}")
    print(f"✓ Successfully loaded: {len(stats['successful'])} ({len(stats['successful'])/total_files*100:.1f}%)")
    print(f"✗ Failed to load: {len(stats['failed'])} ({len(stats['failed'])/total_files*100:.1f}%)")
    
    print("\nBreakdown by file type:")
    for doc_type, counts in stats['total_by_type'].items():
        success_rate = counts['successful']/counts['total']*100 if counts['total'] > 0 else 0
        print(f"  {doc_type:20} {counts['successful']:3}/{counts['total']:3} ({success_rate:5.1f}%)")
    
    if stats['failed']:
        print(f"\n⚠ Failed files ({len(stats['failed'])}):")
        for failed in stats['failed'][:10]:
            print(f"    • {Path(failed['file']).name} ({failed['type']})")
        if len(stats['failed']) > 10:
            print(f"    ... and {len(stats['failed'])-10} more (see ingestion_errors.log)")
        print("\n→ Check 'ingestion_errors.log' for detailed error information")
    
    if not all_docs:
        print("\n❌ No documents were successfully loaded!")
        print("Check 'ingestion_errors.log' for details.")
        return

    print(f"\n✓ Loaded {len(all_docs)} total document chunks")

    import time

    def make_embeddings(embedding_provider, embedding_model, embedding_device, embedding_batch_size, label=""):
        print(f"\nSetting up embeddings for {label}...")
        print(f"  Provider: {embedding_provider}")
        print(f"  Model: {embedding_model}")
        print(f"  Device: {embedding_device}")
        if embedding_provider == "huggingface":
            print("  Using HuggingFace embeddings (supports any model)")
            return HuggingFaceEmbeddings(
                model_name=embedding_model,
                model_kwargs={'device': embedding_device},
                encode_kwargs={'normalize_embeddings': True, 'batch_size': embedding_batch_size}
            )
        else:
            print("  Using FastEmbed embeddings (optimized, limited models)")
            return FastEmbedEmbeddings(
                model_name=embedding_model,
                max_length=512,
                threads=4
            )

    def build_index(chunks, embeddings, index_dir, batch_size, label=""):

        if index_dir.exists():
            shutil.rmtree(index_dir)
            print(f"  Cleared existing index at: {index_dir}")
        index_dir.mkdir(parents=True, exist_ok=True)
        print(f"\nBuilding Chroma index at: {index_dir}")
        print(f"Processing {len(chunks)} chunks in batches of {batch_size}...")

        start_index = time.time()

        if len(chunks) <= batch_size:
            vs = Chroma.from_documents(
                documents=chunks,
                embedding=embeddings,
                persist_directory=str(index_dir),
            )
            del vs
        else:
            vectorstore = Chroma.from_documents(
                documents=chunks[:batch_size],
                embedding=embeddings,
                persist_directory=str(index_dir),
            )
            print(f"  ✓ Batch 1/{(len(chunks) + batch_size - 1) // batch_size} ({batch_size} chunks)")

            for i in range(batch_size, len(chunks), batch_size):
                batch = chunks[i:i + batch_size]
                batch_start = time.time()
                for attempt in range(3):
                    try:
                        vectorstore.add_documents(batch)
                        break
                    except Exception as e:
                        if attempt < 2:
                            wait = 2 ** attempt
                            print(f"  ⚠ Batch upsert failed (attempt {attempt+1}/3), retrying in {wait}s: {e}")
                            time.sleep(wait)
                        else:
                            raise
                batch_num = (i // batch_size) + 1
                total_batches = (len(chunks) + batch_size - 1) // batch_size
                elapsed = time.time() - batch_start
                chunks_per_sec = len(batch) / elapsed if elapsed > 0 else 0
                print(f"  ✓ Batch {batch_num}/{total_batches} ({len(batch)} chunks in {elapsed:.1f}s = {chunks_per_sec:.0f} chunks/s)")

            # Explicitly release the Chroma client to free SQLite connections
            del vectorstore

        print(f"\n✓ {label} index complete in {time.time() - start_index:.2f}s")

    embedding_device = os.environ["EMBEDDING_DEVICE"].lower()
    embedding_provider = os.environ["EMBEDDING_PROVIDER_ML"].lower()
    embedding_model = os.environ["EMBEDDING_MODEL_ML"]
    chunking_strategy = os.environ["CHUNKING_STRATEGY"]
    embedding_batch_size = int(os.environ["EMBEDDING_BATCH_SIZE"])

    embeddings = make_embeddings(embedding_provider, embedding_model, embedding_device, embedding_batch_size, label="Unified (multilingual)")

    print(f"  Splitting into chunks (strategy: {chunking_strategy})...")
    start_split = time.time()
    if chunking_strategy == "adaptive":
        splitter = AdaptiveChunker(embeddings, chunk_size=chunk_size, chunk_overlap=chunk_overlap)
    elif chunking_strategy == "semantic":
        splitter = SemanticTableAwareSplitter(embeddings, chunk_size=chunk_size, chunk_overlap=chunk_overlap)
    else:
        splitter = TableAwareSplitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap)
    chunks = splitter.split_documents(all_docs)
    print(f"  Created {len(chunks)} chunks in {time.time() - start_split:.2f}s")

    chroma_dir = Path(os.environ["CHROMA_DIR"]).resolve()
    build_index(
        chunks,
        embeddings,
        chroma_dir,
        batch_size,
        label="Unified (multilingual)",
    )

    print(f"\n{'='*70}")
    print("FINAL STATISTICS")
    print(f"{'='*70}")
    print(f"  📊 Total documents found:     {total_files}")
    print(f"  ✅ Successfully ingested:     {len(stats['successful'])} ({len(stats['successful'])/total_files*100:.1f}%)")
    print(f"  ❌ Failed to ingest:          {len(stats['failed'])} ({len(stats['failed'])/total_files*100:.1f}%)")
    print(f"  📦 Total chunks created:      {len(chunks):,}")
    print(f"  💾 Unified index:             {chroma_dir}")
    print(f"{'='*70}")
    
    if stats['failed']:
        print(f"\n⚠️  {len(stats['failed'])} files could not be loaded")
        print(f"   See 'ingestion_errors.log' for details")
    else:
        print(f"\n🎉 All files successfully ingested!")
    
    print(f"\n✅ Vector store is ready for querying")

if __name__ == "__main__":
    main()
